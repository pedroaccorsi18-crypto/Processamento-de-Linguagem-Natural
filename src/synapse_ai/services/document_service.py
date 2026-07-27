from __future__ import annotations

import csv
import hashlib
import json
import logging
import re
import zipfile
from collections.abc import Callable
from dataclasses import dataclass
from email import policy
from email.parser import BytesParser
from html import unescape
from io import BytesIO, StringIO
from pathlib import Path
from typing import Protocol
from xml.etree import ElementTree

from synapse_ai.models.document import DocumentStatus

logger = logging.getLogger(__name__)

MAX_UPLOAD_SIZE_BYTES = 10 * 1024 * 1024
SUPPORTED_DOCUMENT_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".pptx",
    ".xlsx",
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".vtt",
    ".eml",
}
SUPPORTED_AUDIO_EXTENSIONS = {".mp3", ".mp4", ".mpeg", ".mpga", ".m4a", ".wav", ".webm", ".ogg"}
TICKET_FIELD_ALIASES = {
    "key": {"key", "issue key", "chave", "ticket", "ticket id", "id"},
    "summary": {"summary", "title", "titulo", "título", "resumo", "assunto"},
    "status": {"status", "situação", "situacao", "estado"},
    "assignee": {"assignee", "assigned to", "responsável", "responsavel", "owner"},
    "reporter": {"reporter", "solicitante", "requester", "criado por", "autor"},
    "priority": {"priority", "prioridade", "severity", "severidade"},
    "issue_type": {"issue type", "tipo", "tipo de item", "tipo de issue"},
    "created": {"created", "created date", "criado", "data de criação", "data de criacao"},
    "updated": {"updated", "updated date", "atualizado", "data de atualização"},
    "due_date": {"due date", "due", "prazo", "vencimento", "data limite"},
    "description": {"description", "descrição", "descricao", "detalhes"},
    "comments": {"comments", "comentários", "comentarios", "comment"},
}
TICKET_FIELD_LABELS = {
    "key": "Chave",
    "summary": "Resumo",
    "status": "Status",
    "assignee": "Responsável",
    "reporter": "Solicitante",
    "priority": "Prioridade",
    "issue_type": "Tipo",
    "created": "Criado em",
    "updated": "Atualizado em",
    "due_date": "Prazo",
    "description": "Descrição",
    "comments": "Comentários",
}
TICKET_FIELD_ORDER = [
    "key",
    "summary",
    "status",
    "priority",
    "issue_type",
    "assignee",
    "reporter",
    "created",
    "updated",
    "due_date",
    "description",
    "comments",
]


class DocumentProcessingError(RuntimeError):
    """Raised when a document cannot be parsed safely."""


@dataclass(frozen=True)
class PlannedDocumentUpload:
    filename: str
    content_type: str
    size_bytes: int


@dataclass(frozen=True)
class UploadedDocument:
    filename: str
    content_type: str
    content: bytes


@dataclass(frozen=True)
class ParsedDocument:
    filename: str
    content_type: str
    size_bytes: int
    text: str
    status: DocumentStatus
    metadata: dict[str, object]


class UploadedFileLike(Protocol):
    name: str
    type: str
    size: int

    def getvalue(self) -> bytes:
        """Return uploaded file bytes."""
        ...


def describe_supported_document_formats() -> list[str]:
    return [
        "PDF",
        "DOCX",
        "PPTX",
        "XLSX",
        "TXT",
        "MD",
        "CSV",
        "JSON",
        "VTT",
        "EML",
        "MP3",
        "M4A",
        "WAV",
    ]


def is_audio_document(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_AUDIO_EXTENSIONS


def validate_planned_upload(upload: PlannedDocumentUpload) -> bool:
    return (
        upload.size_bytes > 0
        and bool(upload.filename.strip())
        and bool(upload.content_type.strip())
    )


def parse_streamlit_upload(uploaded_file: UploadedFileLike) -> ParsedDocument:
    content = uploaded_file.getvalue()
    return parse_uploaded_document(
        UploadedDocument(
            filename=uploaded_file.name,
            content_type=uploaded_file.type or "application/octet-stream",
            content=content,
        )
    )


def parse_uploaded_document(upload: UploadedDocument) -> ParsedDocument:
    extension = Path(upload.filename).suffix.lower()
    _validate_upload(upload, extension, SUPPORTED_DOCUMENT_EXTENSIONS)

    parser = _parser_for_extension(extension)
    try:
        text, parser_metadata = parser(upload.content)
    except DocumentProcessingError:
        raise
    except Exception as exc:  # noqa: BLE001
        logger.warning("Document parsing failed: %s", exc.__class__.__name__)
        raise DocumentProcessingError("Não foi possível extrair texto do arquivo enviado.") from exc

    normalized_text = _normalize_text(text)
    if not normalized_text:
        raise DocumentProcessingError("Nenhum texto foi encontrado no arquivo enviado.")

    metadata: dict[str, object] = {
        "file_extension": extension.removeprefix("."),
        "checksum_sha256": hashlib.sha256(upload.content).hexdigest(),
        "word_count": _word_count(normalized_text),
        "char_count": len(normalized_text),
        **parser_metadata,
    }

    return ParsedDocument(
        filename=upload.filename,
        content_type=upload.content_type,
        size_bytes=len(upload.content),
        text=normalized_text,
        status=DocumentStatus.EXTRACTED,
        metadata=metadata,
    )


def parse_transcribed_audio_document(
    upload: UploadedDocument,
    transcription_text: str,
    transcription_model: str,
) -> ParsedDocument:
    extension = Path(upload.filename).suffix.lower()
    _validate_upload(upload, extension, SUPPORTED_AUDIO_EXTENSIONS)

    normalized_text = _normalize_text(transcription_text)
    if not normalized_text:
        raise DocumentProcessingError("Nenhum texto foi transcrito do áudio enviado.")

    metadata: dict[str, object] = {
        "file_extension": extension.removeprefix("."),
        "checksum_sha256": hashlib.sha256(upload.content).hexdigest(),
        "word_count": _word_count(normalized_text),
        "char_count": len(normalized_text),
        "source_type": "audio_transcription",
        "transcription_model": transcription_model,
    }

    return ParsedDocument(
        filename=upload.filename,
        content_type=upload.content_type,
        size_bytes=len(upload.content),
        text=normalized_text,
        status=DocumentStatus.EXTRACTED,
        metadata=metadata,
    )


def build_document_payload(user_id: str, parsed_document: ParsedDocument) -> dict[str, object]:
    return {
        "user_id": user_id,
        "filename": parsed_document.filename,
        "content_type": parsed_document.content_type,
        "size_bytes": parsed_document.size_bytes,
        "status": parsed_document.status.value,
        "extracted_text": parsed_document.text,
        "text_char_count": len(parsed_document.text),
        "metadata": parsed_document.metadata,
    }


def preview_text(text: str, limit: int = 1200) -> str:
    if len(text) <= limit:
        return text
    return f"{text[:limit].rstrip()}..."


def _validate_upload(
    upload: UploadedDocument,
    extension: str,
    supported_extensions: set[str],
) -> None:
    if not upload.filename.strip():
        raise DocumentProcessingError("O arquivo precisa ter um nome válido.")
    if not upload.content:
        raise DocumentProcessingError("O arquivo enviado está vazio.")
    if len(upload.content) > MAX_UPLOAD_SIZE_BYTES:
        raise DocumentProcessingError("O arquivo excede o limite de 10 MB desta fase.")
    if extension not in supported_extensions:
        raise DocumentProcessingError("Formato ainda não suportado nesta fase.")


def _parser_for_extension(extension: str) -> Callable[[bytes], tuple[str, dict[str, object]]]:
    parsers: dict[str, Callable[[bytes], tuple[str, dict[str, object]]]] = {
        ".txt": _parse_plain_text,
        ".md": _parse_plain_text,
        ".csv": _parse_csv,
        ".json": _parse_json,
        ".vtt": _parse_vtt,
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
        ".pptx": _parse_pptx,
        ".xlsx": _parse_xlsx,
        ".eml": _parse_email,
    }
    return parsers[extension]


def _parse_plain_text(content: bytes) -> tuple[str, dict[str, object]]:
    text, encoding = _decode_text(content)
    return text, {"encoding": encoding}


def _decode_text(content: bytes) -> tuple[str, str]:
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return content.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    raise DocumentProcessingError("Não foi possível identificar a codificação do arquivo de texto.")


def _parse_csv(content: bytes) -> tuple[str, dict[str, object]]:
    text, encoding = _decode_text(content)
    rows = _parse_delimited_rows(text)
    ticket_export = _format_ticket_rows(rows)
    if ticket_export is not None:
        ticket_text, ticket_metadata = ticket_export
        return ticket_text, {"encoding": encoding, **ticket_metadata}
    return text, {"encoding": encoding}


def _parse_json(content: bytes) -> tuple[str, dict[str, object]]:
    text, encoding = _decode_text(content)
    try:
        payload = json.loads(text)
    except json.JSONDecodeError as exc:
        raise DocumentProcessingError("O arquivo JSON enviado não é válido.") from exc

    slack_export = _format_slack_messages(payload)
    if slack_export is not None:
        slack_text, slack_metadata = slack_export
        return slack_text, {"encoding": encoding, **slack_metadata}

    teams_export = _format_teams_messages(payload)
    if teams_export is not None:
        teams_text, teams_metadata = teams_export
        return teams_text, {"encoding": encoding, **teams_metadata}

    return json.dumps(payload, ensure_ascii=False, indent=2), {
        "encoding": encoding,
        "source_type": "json_document",
    }


def _parse_vtt(content: bytes) -> tuple[str, dict[str, object]]:
    text, encoding = _decode_text(content)
    transcript_lines: list[str] = []

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.upper() == "WEBVTT":
            continue
        if "-->" in line:
            continue
        if line.isdigit():
            continue
        if line.startswith(("NOTE", "STYLE", "REGION")):
            continue

        line = re.sub(r"<v\s+([^>]+)>", r"\1: ", line)
        line = line.replace("</v>", "")
        line = _clean_message_text(line)
        if line:
            transcript_lines.append(line)

    return "Transcrição Teams/WebVTT detectada\n\n" + "\n".join(transcript_lines), {
        "encoding": encoding,
        "source_type": "meeting_transcript",
        "transcript_format": "webvtt",
        "line_count": len(transcript_lines),
    }


def _parse_pdf(content: bytes) -> tuple[str, dict[str, object]]:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    page_texts = [(page.extract_text() or "") for page in reader.pages]
    return "\n\n".join(page_texts), {"page_count": len(reader.pages)}


def _parse_docx(content: bytes) -> tuple[str, dict[str, object]]:
    from docx import Document as DocxDocument

    document = DocxDocument(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n\n".join(paragraphs), {"paragraph_count": len(paragraphs)}


def _parse_pptx(content: bytes) -> tuple[str, dict[str, object]]:
    slide_texts: list[str] = []
    with zipfile.ZipFile(BytesIO(content)) as archive:
        slide_paths = sorted(
            path
            for path in archive.namelist()
            if path.startswith("ppt/slides/slide") and path.endswith(".xml")
        )
        for index, slide_path in enumerate(slide_paths, start=1):
            xml_root = ElementTree.fromstring(archive.read(slide_path))
            texts = [
                text_node.text.strip()
                for text_node in xml_root.iter(_xml_name("a", "t"))
                if text_node.text and text_node.text.strip()
            ]
            if texts:
                slide_texts.append(f"Slide {index}\n" + "\n".join(texts))

    return "\n\n".join(slide_texts), {"slide_count": len(slide_paths)}


def _parse_xlsx(content: bytes) -> tuple[str, dict[str, object]]:
    sheet_texts: list[str] = []
    ticket_texts: list[str] = []
    ticket_count = 0
    ticket_fields: set[str] = set()
    row_count = 0
    cell_count = 0
    with zipfile.ZipFile(BytesIO(content)) as archive:
        shared_strings = _read_xlsx_shared_strings(archive)
        sheet_paths = sorted(
            path
            for path in archive.namelist()
            if path.startswith("xl/worksheets/sheet") and path.endswith(".xml")
        )
        for index, sheet_path in enumerate(sheet_paths, start=1):
            rows: list[str] = []
            sheet_rows: list[list[str]] = []
            xml_root = ElementTree.fromstring(archive.read(sheet_path))
            for row in xml_root.iter(_xml_name("main", "row")):
                values = [
                    value
                    for cell in row.iter(_xml_name("main", "c"))
                    if (value := _read_xlsx_cell_value(cell, shared_strings))
                ]
                if values:
                    row_count += 1
                    cell_count += len(values)
                    sheet_rows.append(values)
                    rows.append(" | ".join(values))
            ticket_export = _format_ticket_rows(
                sheet_rows,
                source_label=f"Exportação de tickets detectada - Planilha {index}",
            )
            if ticket_export is not None:
                ticket_text, ticket_metadata = ticket_export
                ticket_texts.append(ticket_text)
                raw_ticket_count = ticket_metadata.get("ticket_count", 0)
                if isinstance(raw_ticket_count, int):
                    ticket_count += raw_ticket_count
                raw_fields = ticket_metadata.get("ticket_fields", [])
                if isinstance(raw_fields, list):
                    ticket_fields.update(str(field) for field in raw_fields)
            if rows:
                sheet_texts.append(f"Planilha {index}\n" + "\n".join(rows))

    if ticket_texts:
        return "\n\n".join(ticket_texts), {
            "sheet_count": len(sheet_paths),
            "row_count": row_count,
            "cell_count": cell_count,
            "source_type": "ticket_export",
            "ticket_count": ticket_count,
            "ticket_fields": sorted(ticket_fields),
        }

    return "\n\n".join(sheet_texts), {
        "sheet_count": len(sheet_paths),
        "row_count": row_count,
        "cell_count": cell_count,
    }


def _parse_email(content: bytes) -> tuple[str, dict[str, object]]:
    message = BytesParser(policy=policy.default).parsebytes(content)
    subject = str(message.get("subject", "") or "")
    sender = str(message.get("from", "") or "")
    recipients = str(message.get("to", "") or "")
    date = str(message.get("date", "") or "")
    parts: list[str] = []

    if message.is_multipart():
        for part in message.walk():
            if part.get_content_maintype() == "multipart":
                continue
            content_type = part.get_content_type()
            if content_type == "text/plain":
                parts.append(str(part.get_content()))
            elif content_type == "text/html" and not parts:
                parts.append(_strip_html(str(part.get_content())))
    else:
        payload = str(message.get_content())
        if message.get_content_type() == "text/html":
            payload = _strip_html(payload)
        parts.append(payload)

    header_lines = [
        f"Assunto: {subject}" if subject else "",
        f"De: {sender}" if sender else "",
        f"Para: {recipients}" if recipients else "",
        f"Data: {date}" if date else "",
    ]
    text = "\n".join(line for line in header_lines if line)
    body = "\n\n".join(part.strip() for part in parts if part.strip())
    if text and body:
        text = f"{text}\n\n{body}"
    elif body:
        text = body

    return text, {
        "subject": subject,
        "from": sender,
        "to": recipients,
        "email_date": date,
        "part_count": len(parts),
    }


def _read_xlsx_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    try:
        root = ElementTree.fromstring(archive.read("xl/sharedStrings.xml"))
    except KeyError:
        return []

    values: list[str] = []
    for item in root.iter(_xml_name("main", "si")):
        texts = [
            text_node.text or ""
            for text_node in item.iter(_xml_name("main", "t"))
            if text_node.text
        ]
        values.append("".join(texts))
    return values


def _read_xlsx_cell_value(
    cell: ElementTree.Element,
    shared_strings: list[str],
) -> str:
    cell_type = cell.attrib.get("t")
    if cell_type == "inlineStr":
        text_node = cell.find(".//" + _xml_name("main", "t"))
        return (text_node.text or "").strip() if text_node is not None else ""

    value_node = cell.find(_xml_name("main", "v"))
    if value_node is None or value_node.text is None:
        return ""

    raw_value = value_node.text.strip()
    if cell_type == "s":
        try:
            return shared_strings[int(raw_value)].strip()
        except (IndexError, ValueError):
            return raw_value
    return raw_value


def _parse_delimited_rows(text: str) -> list[list[str]]:
    first_line = next((line for line in text.splitlines() if line.strip()), "")
    delimiter = max((";", ",", "\t", "|"), key=first_line.count)
    reader = csv.reader(StringIO(text), delimiter=delimiter)
    return [
        [cell.strip() for cell in row]
        for row in reader
        if any(cell.strip() for cell in row)
    ]


def _format_ticket_rows(
    rows: list[list[str]],
    source_label: str = "Exportação de tickets detectada",
) -> tuple[str, dict[str, object]] | None:
    if len(rows) < 2:
        return None

    header_map = _build_ticket_header_map(rows[0])
    if not _looks_like_ticket_export(header_map):
        return None

    lines = [source_label]
    ticket_count = 0
    for row in rows[1:]:
        if not any(cell.strip() for cell in row):
            continue
        ticket_count += 1
        title = _ticket_field(row, header_map, "summary")
        key = _ticket_field(row, header_map, "key")
        heading = title or key or f"Registro {ticket_count}"
        lines.extend(["", f"Ticket {ticket_count}: {heading}"])
        for field in TICKET_FIELD_ORDER:
            value = _ticket_field(row, header_map, field)
            if value:
                lines.append(f"{TICKET_FIELD_LABELS[field]}: {value}")

    if ticket_count == 0:
        return None

    return "\n".join(lines), {
        "source_type": "ticket_export",
        "ticket_count": ticket_count,
        "ticket_fields": sorted(header_map),
    }


def _format_slack_messages(payload: object) -> tuple[str, dict[str, object]] | None:
    messages = [
        message
        for message in _iter_json_message_dicts(payload)
        if _looks_like_slack_message(message)
    ]
    if not messages:
        return None

    lines = ["Exportação Slack detectada"]
    for index, message in enumerate(messages, start=1):
        text = _clean_message_text(_string_value(message.get("text")))
        if not text:
            continue
        author = (
            _string_value(message.get("user"))
            or _string_value(message.get("username"))
            or _string_value(message.get("bot_id"))
            or "autor não identificado"
        )
        timestamp = _string_value(message.get("ts")) or _string_value(message.get("timestamp"))
        subtype = _string_value(message.get("subtype"))

        lines.extend(["", f"Mensagem {index}"])
        lines.append(f"Autor: {author}")
        if timestamp:
            lines.append(f"Data/Hora: {timestamp}")
        if subtype:
            lines.append(f"Tipo: {subtype}")
        lines.append(f"Conteúdo: {text}")

    message_count = len([line for line in lines if line.startswith("Mensagem ")])
    if message_count == 0:
        return None

    return "\n".join(lines), {
        "source_type": "slack_export",
        "platform": "slack",
        "message_count": message_count,
    }


def _format_teams_messages(payload: object) -> tuple[str, dict[str, object]] | None:
    messages = [
        message
        for message in _iter_json_message_dicts(payload)
        if _looks_like_teams_message(message)
    ]
    if not messages:
        return None

    lines = ["Exportação Microsoft Teams detectada"]
    for index, message in enumerate(messages, start=1):
        text = _clean_message_text(_teams_message_text(message))
        if not text:
            continue
        author = _teams_message_author(message) or "autor não identificado"
        timestamp = (
            _string_value(message.get("createdDateTime"))
            or _string_value(message.get("lastModifiedDateTime"))
            or _string_value(message.get("time"))
            or _string_value(message.get("timestamp"))
        )

        lines.extend(["", f"Mensagem {index}"])
        lines.append(f"Autor: {author}")
        if timestamp:
            lines.append(f"Data/Hora: {timestamp}")
        lines.append(f"Conteúdo: {text}")

    message_count = len([line for line in lines if line.startswith("Mensagem ")])
    if message_count == 0:
        return None

    return "\n".join(lines), {
        "source_type": "teams_export",
        "platform": "microsoft_teams",
        "message_count": message_count,
    }


def _iter_json_message_dicts(payload: object) -> list[dict[str, object]]:
    if isinstance(payload, list):
        return [item for item in payload if isinstance(item, dict)]
    if not isinstance(payload, dict):
        return []

    direct_messages: list[dict[str, object]] = []
    for key in ("messages", "items", "value", "chatMessages"):
        value = payload.get(key)
        if isinstance(value, list):
            direct_messages.extend(item for item in value if isinstance(item, dict))

    if direct_messages:
        return direct_messages
    return [payload]


def _looks_like_slack_message(message: dict[str, object]) -> bool:
    return (
        bool(_string_value(message.get("text")))
        and bool({"ts", "user", "username", "bot_id", "subtype"} & set(message))
        and (
            _string_value(message.get("type")) in {"message", ""}
            or "subtype" in message
            or "ts" in message
        )
    )


def _looks_like_teams_message(message: dict[str, object]) -> bool:
    return bool(_teams_message_text(message)) and bool(
        {"createdDateTime", "lastModifiedDateTime", "from", "body", "speaker"} & set(message)
    )


def _teams_message_text(message: dict[str, object]) -> str:
    body = message.get("body")
    if isinstance(body, dict):
        content = body.get("content")
        if isinstance(content, str):
            return content
    for key in ("text", "content", "message", "transcript"):
        value = message.get(key)
        if isinstance(value, str):
            return value
    return ""


def _teams_message_author(message: dict[str, object]) -> str:
    from_value = message.get("from")
    if isinstance(from_value, dict):
        for nested_key in ("user", "application", "device"):
            nested_value = from_value.get(nested_key)
            if isinstance(nested_value, dict):
                display_name = nested_value.get("displayName")
                if isinstance(display_name, str) and display_name.strip():
                    return display_name.strip()
        display_name = from_value.get("displayName")
        if isinstance(display_name, str) and display_name.strip():
            return display_name.strip()

    for key in ("speaker", "author", "sender", "user", "username"):
        value = message.get(key)
        if isinstance(value, str) and value.strip():
            return value.strip()
    return ""


def _clean_message_text(text: str) -> str:
    cleaned = unescape(text)
    cleaned = re.sub(r"<#([^>|]+)\|([^>]+)>", r"#\2", cleaned)
    cleaned = re.sub(r"<#([^>]+)>", r"#\1", cleaned)
    cleaned = re.sub(r"<@([^>]+)>", r"@\1", cleaned)
    cleaned = re.sub(r"<([^>|]+)\|([^>]+)>", r"\2", cleaned)
    cleaned = re.sub(r"<[^>]+>", " ", cleaned)
    return re.sub(r"\s+", " ", cleaned).strip()


def _string_value(value: object) -> str:
    return value.strip() if isinstance(value, str) else ""


def _build_ticket_header_map(header: list[str]) -> dict[str, int]:
    header_map: dict[str, int] = {}
    for index, value in enumerate(header):
        canonical_field = _canonical_ticket_field(value)
        if canonical_field and canonical_field not in header_map:
            header_map[canonical_field] = index
    return header_map


def _canonical_ticket_field(value: str) -> str:
    clean_value = _normalize_header(value)
    for canonical_field, aliases in TICKET_FIELD_ALIASES.items():
        if clean_value in {_normalize_header(alias) for alias in aliases}:
            return canonical_field
    return ""


def _looks_like_ticket_export(header_map: dict[str, int]) -> bool:
    fields = set(header_map)
    return (
        ("summary" in fields or "description" in fields)
        and "status" in fields
        and bool(fields & {"key", "assignee", "priority", "issue_type", "created", "updated"})
    )


def _ticket_field(row: list[str], header_map: dict[str, int], field: str) -> str:
    index = header_map.get(field)
    if index is None or index >= len(row):
        return ""
    return row[index].strip()


def _normalize_header(value: str) -> str:
    return re.sub(r"\s+", " ", value.strip().casefold())


def _xml_name(namespace: str, tag: str) -> str:
    namespaces = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
    }
    return f"{{{namespaces[namespace]}}}{tag}"


def _strip_html(text: str) -> str:
    return re.sub(r"\s+", " ", re.sub(r"<[^>]+>", " ", text)).strip()


def _normalize_text(text: str) -> str:
    return "\n".join(line.rstrip() for line in text.replace("\r\n", "\n").split("\n")).strip()


def _word_count(text: str) -> int:
    return len([word for word in text.split() if word.strip()])
